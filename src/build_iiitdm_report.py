"""
Build docs/report.docx in the IIITDM-SIES internship report
template format:

  Title page -> Certificate -> Abstract -> Acknowledgements ->
  Contents (TOC) -> List of Figures -> List of Tables ->
  Abbreviations -> Symbols ->
  Chapter 1 Introduction (1.1 Background / 1.2 Motivation / 1.3 Objectives) ->
  Chapter 2 Methodology ->
  Chapter 3 Work Done ->
  Chapter 4 Results and Discussions ->
  Chapter 5 Conclusions and Extensions ->
  Bibliography (15 references)

Run:
    python src/build_iiitdm_report.py
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT  = ROOT / "docs" / "report.docx"

BODY_FONT  = "Times New Roman"
CODE_FONT  = "Consolas"
BLACK      = RGBColor(0x00, 0x00, 0x00)

# Personal details
STUDENT_NAME  = "AILA VISHNU VARDHAN"
MENTOR_NAME   = "Dr. Arunachalam Muthiah"
COLLEGE_NAME  = "Vidya Jyothi Institute of Technology"
DEPARTMENT    = "School of Interdisciplinary Design and Innovation (SIDI)"
PROJECT_TITLE = ("A Predictive Machine Learning Framework for Ergonomic Risk "
                 "in Last-Mile Quick-Commerce Delivery Operations")


# ============================================================================
# Low-level helpers
# ============================================================================

def set_font(run, name=BODY_FONT):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)


def add_line(doc, text="", size=12, bold=False, italic=False, align=None,
             space_before=0, space_after=6, font=BODY_FONT, color=BLACK,
             style=None, page_break_before=False):
    """Add a single paragraph.  When `style` is set to a Word heading style
    name ("Heading 1" / "Heading 2" / ...), the paragraph is tagged with
    that style so Word's Navigation Pane, auto-TOC, and PDF outline pick
    it up.  The visual formatting (size / bold / colour) still comes from
    the run so we keep the design consistent regardless of the heading
    style's defaults.

    Set page_break_before=True to attach the page break to the paragraph
    itself instead of adding a separate empty break-paragraph -- avoids
    stray blank pages when the previous chapter/section content ends
    near a page boundary."""
    p = doc.add_paragraph()
    if style is not None:
        p.style = doc.styles[style]
    if page_break_before:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.font.color.rgb = color
        set_font(r, font)
    return p


def add_chapter_heading(doc, chapter_num, title, skip_page_break=False):
    """Chapter heading as a single Heading 2 paragraph 'Chapter N. Title'
    so a STYLEREF field in the running header can pick it up whole.

    Uses paragraph_format.page_break_before instead of an explicit
    doc.add_page_break() so Word does not have to render a separate
    'break paragraph' -- that separate paragraph is what was producing
    the intermittent blank pages between chapters.

    Set skip_page_break=True when a section break has just been added
    (which already starts a new page) to avoid an extra blank page."""
    heading_text = (f"Chapter {chapter_num}. {title}"
                    if chapter_num != "" else title)

    p = doc.add_paragraph(style=doc.styles["Heading 2"])
    if not skip_page_break:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(28)
    r = p.add_run(heading_text)
    r.font.size = Pt(28)
    r.bold = True
    r.font.color.rgb = BLACK
    set_font(r, BODY_FONT)


def add_section_heading(doc, num, title):
    """Main section like '2.1 Two-stage pipeline' -- Heading 4 in Word."""
    p = doc.add_paragraph(style=doc.styles["Heading 4"])
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    r_num = p.add_run(f"{num}   ")
    r_num.font.size = Pt(16)
    r_num.bold = True
    r_num.font.color.rgb = BLACK
    set_font(r_num, BODY_FONT)
    r_title = p.add_run(title)
    r_title.font.size = Pt(16)
    r_title.bold = True
    r_title.font.color.rgb = BLACK
    set_font(r_title, BODY_FONT)


def add_subsection_heading(doc, num, title):
    """Subsection like '2.1.1 ...' -- Heading 5 in Word."""
    p = doc.add_paragraph(style=doc.styles["Heading 5"])
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    r_num = p.add_run(f"{num}   ")
    r_num.font.size = Pt(13)
    r_num.bold = True
    r_num.font.color.rgb = BLACK
    set_font(r_num, BODY_FONT)
    r_title = p.add_run(title)
    r_title.font.size = Pt(13)
    r_title.bold = True
    r_title.font.color.rgb = BLACK
    set_font(r_title, BODY_FONT)


def add_body(doc, text, size=11):
    """Justified body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.4
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.size = Pt(size)
    set_font(r, BODY_FONT)
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(size)
    set_font(r, BODY_FONT)
    return p


def add_numbered(doc, text, size=11):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(size)
    set_font(r, BODY_FONT)
    return p


def add_code_block(doc, lines, size=10):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.6)
        r = p.add_run(line if line else " ")
        r.font.size = Pt(size)
        set_font(r, CODE_FONT)
    add_line(doc, space_before=0, space_after=6)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table(doc, header, rows, header_bg="D9D9D9", widths_cm=None):
    n = len(header)
    t = doc.add_table(rows=1, cols=n)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.font.size = Pt(10)
        r.bold = True
        set_font(r, BODY_FONT)
        set_cell_bg(hdr[i], header_bg)
    for row in rows:
        cells = t.add_row().cells
        for i, txt in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(txt))
            r.font.size = Pt(10)
            set_font(r, BODY_FONT)
    if widths_cm:
        for row in t.rows:
            for i, w in enumerate(widths_cm):
                row.cells[i].width = Cm(w)
    add_line(doc, space_before=0, space_after=6)
    return t


def add_figure(doc, path, caption, width_cm=15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    if Path(path).exists():
        run = p.add_run()
        run.add_picture(str(path), width=Cm(width_cm))
    else:
        r = p.add_run(f"[Figure file missing: {path}]")
        r.italic = True

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    r.font.size = Pt(10)
    r.italic = True
    set_font(r, BODY_FONT)


# ============================================================================
# Page-number section handling: roman for front matter, arabic for body
# ============================================================================

def set_page_number_format(section, fmt="decimal", start=None):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:fmt"), fmt)
    if start is not None:
        pgNumType.set(qn("w:start"), str(start))


def _enable_field_auto_update(doc):
    """Tell Word to refresh all fields (PAGE, STYLEREF, TOC) when the
    document is opened so the running-header chapter name renders
    correctly without the user right-clicking 'Update Field'."""
    settings = doc.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def _add_page_field(run):
    for tag, extras in [("begin", {}), ("instrText", {"text": "PAGE"}),
                        ("separate", {}), ("t", {"text": "1"}),
                        ("end", {})]:
        e = OxmlElement("w:" + ("fldChar" if tag in ("begin", "separate", "end")
                                else tag))
        if tag in ("begin", "separate", "end"):
            e.set(qn("w:fldCharType"), tag)
        if "text" in extras:
            e.text = extras["text"]
        run._element.append(e)


def _add_styleref_field(run, style_name, uppercase=False, placeholder=" "):
    """STYLEREF field: Word auto-fills the run with the most recent
    paragraph text carrying the given style.  Used for chapter names in
    the running header."""
    switches = ' \\* Upper' if uppercase else ''
    parts = [
        ("fldChar", {"type": "begin"}),
        ("instrText", {"text": f'STYLEREF "{style_name}"{switches} \\* MERGEFORMAT'}),
        ("fldChar", {"type": "separate"}),
        ("t", {"text": placeholder}),
        ("fldChar", {"type": "end"}),
    ]
    for tag, attrs in parts:
        e = OxmlElement("w:" + tag)
        if tag == "fldChar":
            e.set(qn("w:fldCharType"), attrs["type"])
        elif "text" in attrs:
            e.text = attrs["text"]
        run._element.append(e)


def _paragraph_border(paragraph, side="bottom", sz=6):
    """Add a hairline border on one side of the paragraph (used to draw
    the thin underline separating running-header from body and the top
    line above the footer)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(sz))
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), "000000")
    pBdr.append(border)


def start_front_matter_section(doc):
    """Insert a section break after the title page and configure the new
    section as the front matter with lowercase Roman page numbers
    (i, ii, iii...) restarting at i, right-aligned in the footer.  The
    title page (section 0 above this break) stays clean with no page
    numeral because we never attach a footer field to it."""
    from docx.enum.section import WD_SECTION
    section = doc.add_section(WD_SECTION.NEW_PAGE)

    footer = section.footer
    footer.is_linked_to_previous = False
    f_p = footer.paragraphs[0]
    f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    r_page = f_p.add_run()
    _add_page_field(r_page)
    r_page.font.size = Pt(10)
    set_font(r_page, BODY_FONT)

    set_page_number_format(section, fmt="lowerRoman", start=1)
    return section


def start_main_body_section(doc):
    """Insert a section break at the current position and configure the
    new section as the main-body running-header layout:
      - header: STYLEREF Heading 2 (chapter name, uppercase), left,
                underlined by a 1px border
      - footer: PAGE field, right, with a 1px top border
      - decimal page numbers restarting at 1
    """
    from docx.enum.section import WD_SECTION
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    # Match the front-matter margins so the text block does not jump width.
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)
    section.different_first_page_header_footer = False

    # -- Header --  chapter name only, left-aligned with a thin underline.
    header = section.header
    header.is_linked_to_previous = False
    h_p = header.paragraphs[0]

    r_left = h_p.add_run()
    _add_styleref_field(r_left, "Heading 2", uppercase=True,
                        placeholder="CHAPTER")
    r_left.font.size = Pt(10)
    r_left.bold = True
    set_font(r_left, BODY_FONT)
    _paragraph_border(h_p, side="bottom")

    # -- Footer --  page number only, right-aligned with a thin overline.
    footer = section.footer
    footer.is_linked_to_previous = False
    f_p = footer.paragraphs[0]
    f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    r_page = f_p.add_run()
    _add_page_field(r_page)
    r_page.font.size = Pt(10)
    r_page.bold = True
    set_font(r_page, BODY_FONT)
    _paragraph_border(f_p, side="top")

    # Restart page numbering at 1 in decimal
    set_page_number_format(section, fmt="decimal", start=1)

    return section


# ============================================================================
# Content assembly
# ============================================================================

def add_title_page(doc):
    # Longer academic title - size 18 keeps it to 3-4 lines cleanly
    add_line(doc, PROJECT_TITLE,
             size=18, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=24, space_after=30)

    add_line(doc, "Internship report submitted under",
             size=12, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=4)
    add_line(doc, "IIITDM-SIES Programme",
             size=13, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=14)
    add_line(doc, "by",
             size=12, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=4)
    add_line(doc, STUDENT_NAME,
             size=14, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=24)

    # Institute logo
    logo_path = ROOT / "assets" / "logos" / "iiitdm_logo.png"
    if logo_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(28)
        r = p.add_run()
        r.add_picture(str(logo_path), width=Cm(6.5))
    else:
        add_line(doc, "[IIITDM Kancheepuram logo]",
                 size=11, italic=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=12, space_after=28,
                 color=RGBColor(0x77, 0x77, 0x77))

    add_line(doc, "SCHOOL OF INTERDISCIPLINARY DESIGN AND INNOVATION (SIDI)",
             size=12, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=3)
    add_line(doc, "INDIAN INSTITUTE OF INFORMATION TECHNOLOGY,",
             size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=1)
    add_line(doc, "DESIGN AND MANUFACTURING, KANCHEEPURAM",
             size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=28)
    add_line(doc, "July 2026",
             size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER)


def add_certificate(doc):
    # Section break from start_front_matter_section already put us on
    # a fresh page (Roman 'i'), so no internal page break is needed --
    # adding one would create a blank page between the title page and
    # Certificate.
    add_line(doc, "Certificate",
             size=28, bold=True, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=24, space_after=36,
             style="Heading 2")

    add_body(doc,
        f"I, {STUDENT_NAME}, from {COLLEGE_NAME} hereby declare that the material "
        f"presented in the Internship Report titled '{PROJECT_TITLE}' represents "
        f"original work carried out by me in the {DEPARTMENT} "
        "at the Indian Institute of Information Technology, Design and "
        "Manufacturing, Kancheepuram during the year 2026. With my signature, I certify "
        "that:")

    for item in [
        "I have not manipulated any of the data or results.",
        "I have not committed any plagiarism of intellectual property. I have clearly "
        "indicated and referenced the contributions of others.",
        "I have explicitly acknowledged all collaborative research and discussions.",
        "I have understood that any false claim will result in severe disciplinary action.",
        "I have understood that the work may be screened for any form of academic "
        "misconduct.",
    ]:
        add_bullet(doc, item)

    add_line(doc, space_before=36, space_after=6)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_ALIGN_PARAGRAPH.RIGHT)
    r1 = p.add_run("Date:")
    r1.font.size = Pt(12)
    set_font(r1, BODY_FONT)
    r2 = p.add_run("\tIntern's Signature")
    r2.font.size = Pt(12)
    set_font(r2, BODY_FONT)

    add_body(doc,
        "In my capacity as supervisor of the above-mentioned work, I certify that the work "
        "presented in this Report is carried out under my supervision, and is worthy of "
        "consideration for the requirements of internship under IIITDM-SIES programme "
        "during the period May 2026 to July 2026.")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_ALIGN_PARAGRAPH.RIGHT)
    r1 = p.add_run(f"Mentor's Name: {MENTOR_NAME}")
    r1.font.size = Pt(12)
    set_font(r1, BODY_FONT)
    r2 = p.add_run("\tMentor's Signature")
    r2.font.size = Pt(12)
    set_font(r2, BODY_FONT)


def add_abstract(doc):
    add_line(doc, "Abstract",
             size=28, bold=True, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=24, space_after=36,
             style="Heading 3", page_break_before=True)

    add_body(doc,
        "Last-mile quick-commerce delivery riders (Blinkit, Zepto) face prolonged "
        "awkward postures, high delivery rates, heavy carrying loads and vehicle "
        "vibration, leading to musculoskeletal disorders (MSDs). Risk is not "
        "uniform across the workforce: age, tenure, daily hours, vehicle type, "
        "and carrying mode all shape the exposure profile, but this variation "
        "stays invisible to the platform. Standard ergonomic methods each cover "
        "part of the picture (RULA and QEC on the observation side, NMQ, "
        "NASA-TLX and Borg CR10 on self-report), and no single tool combines "
        "the six standard ergonomic risk factors into an automated per-rider "
        "screening. This work builds that screening on a cross-sectional sample "
        "of 182 riders from the Chengalpattu and Chennai regions collected in "
        "March and April 2026.")

    add_body(doc,
        "The pipeline runs in two stages. Stage 1 applies standard ergonomic methods "
        "(Borg CR10 action levels, RULA Table C thresholds, and sample-based cuts) to "
        "compute Low, Medium, or High risk labels for each factor from a 182-rider "
        "survey combined with 182 RULA and QEC observation records. Stage 2 trains "
        "supervised classifiers per factor using seven candidate algorithms (Logistic "
        "Regression, Decision Tree, Random Forest, Extra Trees, Hist Gradient Boosting, "
        "XGBoost, and Stacking) tuned via GridSearchCV with SMOTE inside 5-fold "
        "stratified cross-validation.")

    add_body(doc,
        "The five survey-derived factors achieved 58 to 62 percent cross-validation "
        "accuracy with macro AUC between 71 and 76 percent. The Posture model, which "
        "uses the 11 RULA components and 8 QEC scores in addition to the survey "
        "features, reached 97 percent accuracy and 98 percent AUC. Statistical "
        "analysis (chi-square and multivariable logistic regression) identified age, "
        "job duration, workload score, and fatigue score as the strongest individual "
        "predictors of self-reported discomfort. The final deliverable is a Streamlit "
        "web application that accepts the full 36-item questionnaire plus RULA and "
        "QEC observations and returns the six colour-coded risk levels along with "
        "per-factor recommendations.")


def add_acknowledgements(doc):
    add_line(doc, "Acknowledgements",
             size=28, bold=True, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=24, space_after=36,
             style="Heading 3", page_break_before=True)

    add_body(doc,
        "I would like to express my gratitude to my mentor, " + MENTOR_NAME +
        ", who guided me throughout this project. I show deep appreciation "
        "for his support, patience, and technical direction, from the "
        "study design through the modelling and the final web application.")

    add_body(doc,
        f"I would also like to extend my gratitude to the {DEPARTMENT}, "
        "IIITDM Kancheepuram, for hosting the internship and providing the "
        "resources that made this work possible. I would like to thank "
        "Vidya Jyothi Institute of Technology, my home institution, for "
        "the academic foundation that let me take on this project.")

    add_body(doc,
        "I acknowledge the 182 quick-commerce delivery riders in the "
        "Chengalpattu and Chennai regions who participated in the survey "
        "and the observation study. Without their time and honesty this "
        "work would not have been possible.")

    add_body(doc,
        "Finally, I would like to thank my friends and family for keeping "
        "me motivated throughout the internship.")


def add_toc_placeholder(doc):
    """Hand-populated table of contents in the same visual style as List
    of Figures / List of Tables (tab stops + dot leader).  Page numbers
    are anchored to the LoF entries so they stay in sync with the actual
    layout, and are still refreshable via Word's "Update Field" if the
    user ever inserts an auto-TOC."""
    add_line(doc, "Contents",
             size=28, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=24, space_after=24,
             style="Heading 1", page_break_before=True)

    # Page anchors below match what Word actually renders (verified by
    # driving Word COM to convert the built docx to PDF and reading the
    # PAGE field on every physical page).  Front-matter Roman numerals
    # start at ii because the title page is unnumbered per the
    # different-first-page footer flag.  Main-body decimals restart at
    # 1 from Chapter 1.
    contents = [
        ("Certificate",                                            "i"),
        ("Abstract",                                               "ii"),
        ("Acknowledgements",                                       "iii"),
        ("Contents",                                               "iv"),
        ("List of Figures",                                        "v"),
        ("List of Tables",                                         "vi"),
        ("Abbreviations",                                          "vii"),
        ("Symbols",                                                "viii"),

        ("1   INTRODUCTION",                                       "1"),
        ("    1.1   Background",                                   "1"),
        ("    1.2   Motivation",                                   "1"),
        ("    1.3   Objectives of the work",                       "2"),

        ("2   METHODOLOGY",                                        "3"),
        ("    2.1   Two-stage design",                             "3"),
        ("    2.2   Study design and data sources",                "3"),
        ("        2.2.1   Rider survey",                           "4"),
        ("        2.2.2   Posture observations",                   "4"),
        ("    2.3   Data cleaning and feature engineering",        "4"),
        ("    2.4   Stage 1: Deterministic risk scoring",          "5"),
        ("    2.5   Statistical analysis",                         "5"),
        ("    2.6   Stage 2: Machine learning pipeline",           "6"),
        ("        2.6.1   Logistic Regression",                    "6"),
        ("        2.6.2   Decision Tree",                          "6"),
        ("        2.6.3   Random Forest",                          "6"),
        ("        2.6.4   Extra Trees",                            "6"),
        ("        2.6.5   Histogram Gradient Boosting",            "7"),
        ("        2.6.6   XGBoost",                                "7"),
        ("        2.6.7   Stacking Classifier",                    "7"),
        ("        2.6.8   SMOTE and cross-validation",             "7"),
        ("        2.6.9   Per-target feature exclusions",          "7"),
        ("    2.7   Evaluation",                                   "7"),

        ("3   WORK DONE",                                          "8"),
        ("    3.1   Notebook pipeline",                            "8"),
        ("    3.2   Instrument design and preparation",            "9"),
        ("    3.3   Data collection execution",                    "9"),
        ("    3.4   Data cleaning, encoding, and integration",     "10"),
        ("    3.5   Statistical analysis and model training runs", "11"),
        ("    3.6   Web application development and deployment",   "11"),
        ("    3.7   Iterations and fixes",                         "12"),
        ("        3.7.1   Repetition binning correction",          "12"),
        ("        3.7.2   Duration leakage correction",            "12"),
        ("        3.7.3   Posture model with RULA and QEC inputs", "12"),

        ("4   RESULTS AND DISCUSSION",                             "13"),
        ("    4.1   Sample profile",                               "13"),
        ("    4.2   NMQ pain prevalence and statistical predictors","15"),
        ("    4.3   Stage-1 risk distribution",                    "17"),
        ("    4.4   Stage-2 model performance",                    "19"),
        ("    4.5   Per-class metrics and feature importance",     "21"),
        ("    4.6   Web application demonstration",                "22"),

        ("5   CONCLUSIONS AND EXTENSIONS",                         "25"),

        ("BIBLIOGRAPHY",                                           "27"),
    ]

    # A row is a "chapter row" (bold) if the label starts at column 0
    # and the leading token is either a bare digit (main-body chapter)
    # or the literal word BIBLIOGRAPHY.
    def _is_chapter_row(lbl):
        stripped = lbl.lstrip()
        if stripped != lbl:            # subsection / section row
            return False
        return (stripped == "BIBLIOGRAPHY"
                or (stripped[:1].isdigit() and stripped[:2].isupper()
                    is False and " " in stripped))

    for label, page in contents:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Cm(16.0), WD_ALIGN_PARAGRAPH.RIGHT)
        r = p.add_run(f"{label}\t{page}")
        r.font.size = Pt(11)
        set_font(r, BODY_FONT)
        if _is_chapter_row(label):
            r.font.bold = True
        pPr = p._p.get_or_add_pPr()
        tabs_el = pPr.find(qn("w:tabs"))
        if tabs_el is not None:
            last_tab = tabs_el.findall(qn("w:tab"))[-1]
            last_tab.set(qn("w:leader"), "dot")


def add_list_of_figures(doc):
    add_line(doc, "List of Figures",
             size=28, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=24, space_after=24,
             style="Heading 1", page_break_before=True)
    # Page anchors are aligned to Chapter 4 sections' actual pages
    # in the rendered Word output (Chapter 4 spans main-body 13-24).
    figures = [
        ("2.1",  "Pipeline overview: raw inputs to interactive prediction.",                  "3"),
        ("4.1",  "Sample profile: age, platform, vehicle, and carrying mode.",                "14"),
        ("4.2",  "NMQ 12-month pain prevalence per body area.",                                "15"),
        ("4.3",  "Discomfort prevalence broken down by demographic group.",                    "16"),
        ("4.4",  "Stage-1 Low / Medium / High counts per risk factor.",                        "17"),
        ("4.5",  "Discomfort prevalence within each Low / Medium / High band.",                "18"),
        ("4.6",  "Pearson correlation matrix across the numeric feature pool.",                "18"),
        ("4.7",  "Confusion matrices for the best model per factor.",                          "19"),
        ("4.8",  "ROC curves (one-vs-rest) for the best model per factor.",                    "20"),
        ("4.9",  "Top 10 features by importance for the best model per factor.",               "21"),
        ("4.10", "Web app: sample-profile shortcuts and demographic section.",                 "22"),
        ("4.11", "Web app: Nordic Musculoskeletal Questionnaire section.",                     "23"),
        ("4.12", "Web app: NASA-TLX and Borg CR10 sliders.",                                   "23"),
        ("4.13", "Web app: RULA and QEC observation sections.",                                "24"),
        ("4.14", "Web app: predicted risk profile output.",                                    "24"),
    ]
    for num, cap, page in figures:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        # left tab for caption at 3 cm, right tab with dot leader for page at 16 cm
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Cm(3.0))
        tab = tabs.add_tab_stop(Cm(16.0), WD_ALIGN_PARAGRAPH.RIGHT)
        # Add dot leader through XML (python-docx doesn't expose it directly)
        _ = tab
        r = p.add_run(f"Figure {num}\t{cap}\t{page}")
        r.font.size = Pt(11)
        set_font(r, BODY_FONT)
        # Set dot leader on the right tab via XML
        pPr = p._p.get_or_add_pPr()
        tabs_el = pPr.find(qn("w:tabs"))
        if tabs_el is not None:
            last_tab = tabs_el.findall(qn("w:tab"))[-1]
            last_tab.set(qn("w:leader"), "dot")


def add_list_of_tables(doc):
    add_line(doc, "List of Tables",
             size=28, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=24, space_after=24,
             style="Heading 1", page_break_before=True)
    tables = [
        ("2.1", "Per-target feature exclusions to prevent label leakage.",                    "7"),
        ("4.1", "NMQ 12-month pain prevalence per body area.",                                "15"),
        ("4.2", "Chi-square test: risk factor vs self-reported discomfort.",                  "15"),
        ("4.3", "Significant predictors of discomfort from multivariable "
                "logistic regression.",                                                        "16"),
        ("4.4", "Stage-1 risk band counts per factor.",                                        "17"),
        ("4.5", "Best Stage-2 model per risk factor: 5-fold stratified CV.",                   "19"),
        ("4.6", "Per-class ROC AUC (one-vs-rest) for the best model per factor.",              "20"),
        ("4.7", "Per-class precision, recall, F1, and support.",                               "20"),
        ("4.8", "Top 5 most important features per factor.",                                   "21"),
        ("4.9", "Winning hyperparameters per target after GridSearchCV.",                      "21"),
    ]
    for num, cap, page in tables:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Cm(3.0))
        tabs.add_tab_stop(Cm(16.0), WD_ALIGN_PARAGRAPH.RIGHT)
        r = p.add_run(f"Table {num}\t{cap}\t{page}")
        r.font.size = Pt(11)
        set_font(r, BODY_FONT)
        pPr = p._p.get_or_add_pPr()
        tabs_el = pPr.find(qn("w:tabs"))
        if tabs_el is not None:
            last_tab = tabs_el.findall(qn("w:tab"))[-1]
            last_tab.set(qn("w:leader"), "dot")


def add_abbreviations(doc):
    add_line(doc, "Abbreviations",
             size=28, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=24, space_after=24,
             style="Heading 1", page_break_before=True)
    abbrevs = [
        ("MSD",       "Musculoskeletal Disorder"),
        ("NMQ",       "Nordic Musculoskeletal Questionnaire"),
        ("RULA",      "Rapid Upper Limb Assessment"),
        ("QEC",       "Quick Exposure Check"),
        ("NASA-TLX",  "NASA Task Load Index"),
        ("CV",        "Cross-Validation"),
        ("SMOTE",     "Synthetic Minority Over-sampling Technique"),
        ("ML",        "Machine Learning"),
        ("CSV",       "Comma-Separated Values"),
        ("ROC",       "Receiver Operating Characteristic"),
        ("AUC",       "Area Under the Curve"),
        ("OR",        "Odds Ratio"),
        ("CI",        "Confidence Interval"),
        ("HistGBM",   "Histogram-based Gradient Boosting Machine"),
        ("XGBoost",   "eXtreme Gradient Boosting"),
        ("CART",      "Classification and Regression Trees"),
        ("dph",       "Deliveries Per Hour"),
        ("API",       "Application Programming Interface"),
    ]
    for abbr, full in abbrevs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(3.5))
        r1 = p.add_run(abbr)
        r1.font.size = Pt(11)
        r1.bold = True
        set_font(r1, BODY_FONT)
        r2 = p.add_run(f"\t{full}")
        r2.font.size = Pt(11)
        set_font(r2, BODY_FONT)


def add_symbols(doc):
    add_line(doc, "Symbols",
             size=28, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=24, space_after=24,
             style="Heading 1", page_break_before=True)
    symbols = [
        ("n",   "sample size (n = 182 riders)"),
        ("k",   "number of nearest minority neighbours used by SMOTE"),
        ("λ",   "uniform random weight for SMOTE interpolation"),
        ("β",   "regression coefficient in logistic regression"),
        ("χ²",  "chi-square test statistic"),
        ("p",   "p-value (statistical significance)"),
        ("dph", "deliveries per hour, deliveries_num / work_hours_num"),
    ]
    for sym, meaning in symbols:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(3.5))
        r1 = p.add_run(sym)
        r1.font.size = Pt(11)
        r1.italic = True
        r1.bold = True
        set_font(r1, BODY_FONT)
        r2 = p.add_run(f"\t{meaning}")
        r2.font.size = Pt(11)
        set_font(r2, BODY_FONT)


# ============================================================================
# Chapters
# ============================================================================

def chapter_1(doc, skip_page_break=False):
    add_chapter_heading(doc, 1, "Introduction",
                        skip_page_break=skip_page_break)

    add_section_heading(doc, "1.1", "Background")
    add_body(doc,
        "Last-mile quick-commerce delivery riders for platforms like Blinkit and Zepto spend long hours "
        "on bikes, climb stairs, carry packages, and ride in heavy traffic. Over "
        "time this kind of work causes musculoskeletal disorders (MSDs): back pain, "
        "shoulder pain, wrist injuries, knee problems. The risk is not the same "
        "for every rider. Someone who is 22 and works 4 hours a day on a scooter is "
        "in a very different situation from a 45-year-old who works 10 hours a day "
        "on a motorbike with a handheld bag.")
    add_body(doc,
        "Standard ergonomic methods cover different facets of this problem. RULA "
        "and QEC handle the observational side (posture, body-region exposure), "
        "while NMQ, NASA-TLX, and Borg CR10 cover what the rider reports about pain, "
        "mental load, and perceived exertion. None of them on its own gives a "
        "complete picture. The project pulls all of them together into one "
        "per-rider profile across six ergonomic dimensions: Force, Repetition, "
        "Posture, Duration, Contact Stress, Vibration.")

    add_section_heading(doc, "1.2", "Motivation")
    add_body(doc,
        "The delivery-platform economy in India has grown rapidly in recent years. "
        "Riders form the largest occupational group in this economy and face "
        "physical demands that are qualitatively different from traditional "
        "warehouse or manufacturing work: long distances on two-wheelers, unstable "
        "load positioning, unpredictable traffic, weather exposure, and constant "
        "algorithmic time pressure. There is no automated tool that a platform can "
        "use to screen individual riders across all six standard ergonomic risk "
        "factors from a single questionnaire.")
    add_body(doc,
        "A per-rider risk profile would let a platform intervene early: shorten "
        "shifts for high-risk profiles, provide equipment reviews, offer targeted "
        "physiotherapy, and re-screen riders periodically. Without such a tool, "
        "risk stays invisible until MSDs surface as absenteeism or medical "
        "consultations. This project addresses that gap by combining published "
        "ergonomic assessment methods with modern classifiers into a single "
        "screening pipeline and web application.")

    add_section_heading(doc, "1.3", "Objectives of the work")
    add_numbered(doc,
        "Design a two-stage pipeline that combines a survey of quick-commerce delivery "
        "riders with standard ergonomic observation tools (RULA and QEC) into "
        "per-rider labels for six risk factors: Force, Repetition, Posture, "
        "Duration, Contact Stress, and Vibration.")
    add_numbered(doc,
        "Implement the pipeline reproducibly in Python (Jupyter notebooks, "
        "scikit-learn, XGBoost, imbalanced-learn, statsmodels) with a fixed random "
        "seed and end-to-end regeneration of every intermediate artefact.")
    add_numbered(doc,
        "Train and evaluate seven candidate ML algorithms per risk factor using "
        "5-fold stratified cross-validation with SMOTE oversampling inside every "
        "fold, and report accuracy, macro F1, and macro AUC.")
    add_numbered(doc,
        "Identify and correct methodological issues discovered during modelling "
        "(class boundary degeneracy in the Repetition binning, feature leakage in "
        "the Duration model) and document each fix openly.")
    add_numbered(doc,
        "Deploy the trained models behind a Streamlit web application that "
        "accepts the full 36-item questionnaire plus the 11 RULA components and "
        "8 QEC scores, and returns the six colour-coded risk levels with "
        "per-factor recommendations.")


def chapter_2(doc):
    add_chapter_heading(doc, 2, "Methodology")

    add_section_heading(doc, "2.1", "Two-stage design")
    add_body(doc,
        "The pipeline runs in two stages. Stage 1 computes Low, Medium, or High "
        "labels for each of the six ergonomic risk factors using deterministic "
        "rules from standard ergonomic methods: Borg CR10 action levels for Force, "
        "RULA Table C action levels for Posture, sample terciles or fixed cuts for "
        "the remaining factors. Stage 1 is the ground truth that Stage 2 tries to "
        "learn.")
    add_body(doc,
        "Stage 2 trains a supervised classifier per risk factor. Per-target "
        "feature exclusions remove the inputs that define the label so the model "
        "has to learn from the remaining rider profile rather than memorising the "
        "Stage-1 rule. Each classifier is evaluated by 5-fold stratified "
        "cross-validation with SMOTE oversampling inside every training fold.")
    add_body(doc,
        "The two-stage split lets an ergonomist audit the Stage-1 labels by hand "
        "and a machine-learning reviewer audit the Stage-2 classifiers "
        "independently, without either review interfering with the other. The "
        "trained classifiers can then be deployed behind a web form that runs on "
        "a new rider's profile without re-applying the RULA and QEC worksheets.")

    add_figure(doc, ROOT / "outputs" / "figures" / "methodology_flowchart.png",
               "Figure 2.1: Pipeline overview from raw inputs to interactive "
               "prediction. Rider survey and posture observations are cleaned, "
               "engineered, then labelled by the Stage-1 rules; Stage 2 learns "
               "the labels from the profile.")

    add_section_heading(doc, "2.2", "Study design and data sources")

    add_body(doc,
        "This is a cross-sectional study conducted in the Chengalpattu and Chennai "
        "regions to evaluate the work conditions and posture among the quick "
        "delivery services which are Blinkit and Zepto. The data was collected "
        "in March and April of 2026. A convenience sampling method was used to "
        "approach delivery partners and collect the data.")

    add_body(doc,
        "The same 182 participants supplied both data streams described below. "
        "Rider self-report was captured through a Google Form questionnaire; "
        "posture was recorded on the same rider during a delivery task by a "
        "trained observer. Study eligibility was applied at the person level "
        "before either instrument was administered.")

    add_line(doc, "Inclusion criteria",
             size=12, bold=True, space_before=6, space_after=4)
    add_bullet(doc,
        "Full-time delivery partners working on either the Blinkit or the "
        "Zepto platform.")
    add_bullet(doc,
        "Minimum three months of continuous delivery experience.")
    add_bullet(doc,
        "Using a two-wheeler (scooter or motorbike) for delivery.")
    add_bullet(doc,
        "Both men and women, aged between 18 and 55 years.")

    add_line(doc, "Exclusion criteria",
             size=12, bold=True, space_before=6, space_after=4)
    add_bullet(doc,
        "Chronic musculoskeletal disorders that developed before the start "
        "of delivery work.")
    add_bullet(doc,
        "Accidental surgery within the past three months.")

    add_subsection_heading(doc, "2.2.1", "Rider survey")
    add_body(doc,
        "The file delivery_rider_survey.csv contains 182 self-administered "
        "responses to a 36-item questionnaire covering demographics, lifestyle, "
        "work pattern, Nordic Musculoskeletal Questionnaire 12-month pain across "
        "9 body areas, 7-day discomfort in 4 areas, five discomfort-related "
        "follow-ups, six NASA-TLX workload items (each on a 0-100 scale), and "
        "six Borg CR10 perceived-exertion items (each on a 0-10 scale). After "
        "cleaning the file contains 182 rows and 48 columns.")

    add_subsection_heading(doc, "2.2.2", "Posture observations")
    add_body(doc,
        "The file posture_data.xlsx contains 182 ergonomic observation records "
        "across the workbook's sheets. Each record carries 11 RULA components "
        "(upper arm, lower arm, wrist, wrist twist, neck position, trunk "
        "position, legs, muscle A, force A, muscle B, force B), 3 derived RULA "
        "Table scores, and 8 QEC scores. The observation rows do not carry rider "
        "identifiers; the pairing to the survey uses the severity-rank merge "
        "described in section 2.3.")

    add_section_heading(doc, "2.3", "Data cleaning and feature engineering")
    add_body(doc,
        "Phase 1 (01_data_cleaning.ipynb) normalises the raw CSV, standardises "
        "Yes/No answers, and validates row counts. Phase 2 "
        "(02_feature_engineering.ipynb) encodes ordered categorical bins into "
        "integer codes, assigns numeric midpoints to bin ranges that will feed "
        "arithmetic formulae, ranks vehicle type and carrying mode, creates "
        "short-named binary aliases for the 18 Yes/No survey items, and derives "
        "composite scores.")
    add_body(doc,
        "Four composite features are derived: workload_score (mean of the six "
        "NASA-TLX items with satisfaction reversed), fatigue_score (mean of the "
        "six Borg CR10 items), force_exertion (Borg lifting item as its own "
        "feature), and vibration_index (vehicle_rank multiplied by "
        "work_hours_num). Five product-form interactions are also added so "
        "tree-based models can pick up combined effects such as an older rider "
        "with high workload.")
    add_body(doc,
        "The posture observations, which lack rider identifiers, are merged into "
        "the survey via a severity-rank merge. Each rider receives an "
        "exposure_severity score from their normalised NMQ pain count, fatigue "
        "score, and working hours. Riders are ranked by this severity descending; "
        "posture rows are ranked by RULA Table C descending; the two ranked lists "
        "are joined one-to-one. The methodological caveat is discussed in the "
        "Limitations section of Chapter 5.")

    add_section_heading(doc, "2.4", "Stage 1: Deterministic risk scoring")
    add_body(doc,
        "Phase 3 (03_risk_scoring.ipynb) applies the following rules and writes "
        "risk_profile.csv:")
    add_bullet(doc, "Force: Borg CR10 action levels applied to force_exertion "
                    "(0-3 Low, 4-6 Medium, 7-10 High).")
    add_bullet(doc, "Repetition: deliveries per hour with fixed cuts at 2.5 and "
                    "3.75. The earlier version used pandas qcut into terciles, "
                    "which had a boundary degeneracy documented in Chapter 3.")
    add_bullet(doc, "Duration: continuous riding hours (5 or fewer Low, 6-7 "
                    "Medium, over 7 High).")
    add_bullet(doc, "Vibration: vibration_index binned by sample tercile.")
    add_bullet(doc, "Contact Stress: composite of carrying_contact_rank, "
                    "work_hours_num, and a small age multiplier, binned by "
                    "sample tercile.")
    add_bullet(doc, "Posture: RULA Table C action levels (1-2 Low, 3-4 Medium, "
                    "5 or higher High). The training data minimum is 3, so the "
                    "Posture model in practice has no Low class.")

    add_section_heading(doc, "2.5", "Statistical analysis")
    add_body(doc,
        "Phase 5 (05_stats.ipynb) runs two analyses on the outcome variable "
        "discomfort (binary: 1 if the rider reported pain in any NMQ 12-month "
        "area). First, a chi-square test of independence "
        "(scipy.stats.chi2_contingency) tests whether each risk factor is "
        "associated with discomfort at p less than 0.05. Second, a multivariable "
        "logistic regression (statsmodels.api.Logit) estimates the effect of "
        "eleven predictors while controlling for the others; results are reported "
        "as odds ratios with 95 percent confidence intervals.")

    add_section_heading(doc, "2.6", "Stage 2: Machine learning pipeline")
    add_body(doc,
        "Phase 6 (06_modeling.ipynb) trains a classifier per target using "
        "seven candidate algorithms. Each algorithm sits inside an "
        "imblearn.Pipeline with SMOTE oversampling upstream so the "
        "resampling only affects training folds. Hyperparameters are tuned "
        "with GridSearchCV on the macro-F1 score inside 5-fold stratified "
        "cross-validation, and the best model per target is refit on the "
        "full sample and saved as outputs/models/best_<factor>.pkl.")

    add_subsection_heading(doc, "2.6.1", "Logistic Regression")
    add_body(doc,
        "Regularised linear classifier used as the baseline. Runs with L2 "
        "penalty and class_weight='balanced' so the minority Low / Medium "
        "class does not get ignored on imbalanced targets. Serves as a "
        "floor: any tree-based method that fails to beat Logistic "
        "Regression on a factor is either overfitting or getting nothing "
        "from feature interactions.")

    add_subsection_heading(doc, "2.6.2", "Decision Tree")
    add_body(doc,
        "Single CART decision tree with tuned max_depth (grid 3 to 10) "
        "and min_samples_leaf (grid 2 to 10). Included to establish the "
        "value added by ensembling; not competitive on its own for any "
        "of the six targets.")

    add_subsection_heading(doc, "2.6.3", "Random Forest")
    add_body(doc,
        "Bagged ensemble of 300 to 500 trees with bootstrap sampling and "
        "sqrt(p) features per split. Its low-variance predictions and "
        "graceful handling of mixed feature types (encoded categoricals "
        "plus continuous scores) made it the winning model on three of "
        "the six targets: Repetition, Duration, and Contact Stress.")

    add_subsection_heading(doc, "2.6.4", "Extra Trees")
    add_body(doc,
        "Extremely randomised trees: like Random Forest but with fully "
        "random splits at every node (no split-quality optimisation on "
        "features). The extra randomisation trades a small drop in "
        "individual-tree accuracy for lower between-tree correlation, "
        "which helped it win on Vibration.")

    add_subsection_heading(doc, "2.6.5", "Histogram Gradient Boosting")
    add_body(doc,
        "Sequential ensemble that fits new trees to the residuals of "
        "the current model, using histogram binning of the input features "
        "for training speed. Won on Force and Posture, the two targets "
        "where the training signal is either strongest (Posture, thanks "
        "to the RULA and QEC observation features) or most nonlinear "
        "(Force, with the Borg CR10 cutoffs).")

    add_subsection_heading(doc, "2.6.6", "XGBoost")
    add_body(doc,
        "Gradient-boosted trees with L1 and L2 regularisation on the "
        "leaf weights and column subsampling per tree. Placed second on "
        "several factors but did not win any; the tuned Random Forest "
        "and HistGBM absorbed the same signal at the default search "
        "budget.")

    add_subsection_heading(doc, "2.6.7", "Stacking Classifier")
    add_body(doc,
        "Meta-learner that takes the out-of-fold predictions of the "
        "Random Forest, Extra Trees, HistGBM, and XGBoost base learners "
        "as inputs and fits a Logistic Regression on top. Failed to "
        "beat the strongest base learner on any factor because the four "
        "base learners largely agreed with each other; stacking helps "
        "when base learners disagree, and here they did not.")

    add_subsection_heading(doc, "2.6.8", "SMOTE and cross-validation")
    add_body(doc,
        "SMOTE (Synthetic Minority Over-sampling Technique) synthesises "
        "new minority-class samples along the segments joining a "
        "minority point to its k nearest minority neighbours. Placing it "
        "inside imblearn.Pipeline ensures the synthesis runs only on "
        "training folds during cross-validation, so validation-fold "
        "accuracy is never inflated by peeking at test-fold neighbours. "
        "k_neighbors is set per target so it never exceeds the minority "
        "training-fold count. GridSearchCV picks the best hyperparameters "
        "per algorithm-target pair on macro-F1 (StratifiedKFold with "
        "shuffle=True, random_state=42).")

    add_subsection_heading(doc, "2.6.9", "Per-target feature exclusions")
    add_body(doc,
        "Each target excludes the survey inputs that directly define its "
        "Stage-1 label; otherwise the model would just memorise the rule "
        "(label leakage). The Posture target also receives the 11 RULA "
        "components and 8 QEC scores as extra features, giving 63 total.")
    add_table(doc,
              header=["Target", "Excluded features", "Final feature count"],
              rows=[
                  ["Force", "force_exertion, force_x_age", "42"],
                  ["Repetition", "deliveries_num, work_hours_num, deliv_x_days", "41"],
                  ["Duration", "work_hours_num, vibration_index", "42"],
                  ["Vibration", "vibration_index, vehicle_rank, work_hours_num", "41"],
                  ["Contact Stress", "carrying_contact_rank, work_hours_num", "42"],
                  ["Posture", "(none from survey pool)", "44 + 11 + 8 = 63"],
              ],
              widths_cm=[3.5, 8.0, 4.0])
    add_line(doc, "Table 2.1: Per-target feature exclusions to prevent label "
                  "leakage.",
             size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=12)

    add_section_heading(doc, "2.7", "Evaluation")
    add_body(doc,
        "Phase 7 (07_evaluation.ipynb) uses cross_val_predict to produce "
        "out-of-fold predictions for every rider. The evaluation outputs are: "
        "per-class precision, recall, F1, and support from "
        "classification_report; confusion matrices per target; one-vs-rest ROC "
        "curves and per-class AUC (macro AUC is reported as the unweighted "
        "mean); and feature importance (feature_importances_ for tree models, "
        "abs(coef_) for Logistic Regression).")


def chapter_3(doc):
    add_chapter_heading(doc, 3, "Work Done")

    add_section_heading(doc, "3.1", "Notebook pipeline")
    add_body(doc,
        "The full pipeline is implemented across seven Jupyter notebooks "
        "(01 through 07) plus a Streamlit application. Every notebook is "
        "designed to run standalone from its previous input file, so the whole "
        "pipeline can be regenerated end to end with a fixed random seed of 42.")
    add_bullet(doc, "01_data_cleaning normalises the raw CSV and validates row "
                    "counts.")
    add_bullet(doc, "02_feature_engineering encodes categoricals, derives "
                    "composite scores, aliases the 18 binary items, and merges "
                    "the posture observations by severity rank.")
    add_bullet(doc, "03_risk_scoring applies standard thresholds to compute "
                    "Low/Medium/High labels per risk factor.")
    add_bullet(doc, "04_eda produces demographics, NMQ prevalence, risk "
                    "distribution, discomfort-by-group, and correlation heatmap "
                    "figures.")
    add_bullet(doc, "05_stats runs chi-square and multivariable logistic "
                    "regression.")
    add_bullet(doc, "06_modeling runs SMOTE plus GridSearchCV across 7 "
                    "algorithms per target and keeps the best by F1 macro.")
    add_bullet(doc, "07_evaluation produces confusion matrices, ROC curves, "
                    "classification reports, and feature importance plots.")

    add_section_heading(doc, "3.2", "Instrument design and preparation")
    add_body(doc,
        "Two data-collection instruments were prepared before fieldwork "
        "began. The self-report instrument was a 36-item Google Form "
        "structured in five modules: demographics and work pattern "
        "(seventeen items), the Nordic Musculoskeletal Questionnaire "
        "covering 12-month pain across nine body areas (Q18) and 7-day "
        "discomfort across four areas (Q19), five discomfort-related "
        "follow-up items on performance impact, leave taken, medical "
        "consultation, and worsening triggers (Q20 to Q24), six NASA-TLX "
        "workload items on 0 to 100 sliders (Q25 to Q30), and six Borg "
        "CR10 perceived-exertion items on 0 to 10 sliders (Q31 to Q36). "
        "The 'dissatisfied with own performance' NASA-TLX item was written "
        "as a positive-polarity slider so all six workload items point the "
        "same direction; the polarity is reversed at analysis time.")
    add_body(doc,
        "The observation instrument was a paper form carrying the eleven "
        "standard RULA elements (upper arm, lower arm, wrist, wrist twist, "
        "neck position, trunk position, legs, muscle A, force A, muscle B, "
        "force B), the three RULA Table lookup scores (A, B, C), and the "
        "eight QEC region and exposure scores (back, shoulder / arm, "
        "wrist / hand, neck, driving, vibration, work pace, stress). Both "
        "instruments were piloted with three delivery partners outside the "
        "study region so the item wording, response options, and observer "
        "worksheet layout could be revised before real fieldwork began.")

    add_section_heading(doc, "3.3", "Data collection execution")
    add_body(doc,
        "Fieldwork ran across the Chengalpattu and Chennai regions between "
        "March and April 2026, following the eligibility criteria and "
        "instruments described in Section 2.2. Recruitment used convenience "
        "sampling at platform handover points where delivery partners "
        "regularly cluster between shifts.")
    add_body(doc,
        "Each participant gave verbal consent before either instrument "
        "was administered. The 36-item Google Form was completed on a "
        "researcher-provided phone in about 15 to 20 minutes. Posture was "
        "recorded on paper by a trained observer during a natural delivery "
        "task, capturing the 11 RULA elements and 8 QEC region / exposure "
        "scores in one pass. Riders whose demographic or NMQ responses "
        "were incomplete after follow-up were excluded and are not part "
        "of the final 182. Data was anonymised at the point of collection; "
        "no rider identifiers appear in the released model_ready.csv or "
        "posture_data.xlsx files.")

    add_section_heading(doc, "3.4",
                        "Data cleaning, encoding, and integration")
    add_body(doc,
        "Data preparation ran in two notebooks. 01_data_cleaning loaded "
        "the raw Google Form CSV, standardised text answers to canonical "
        "Yes / No values, corrected a handful of column-name typos, and "
        "dropped rows that were incomplete after follow-up. The cleaned "
        "CSV was 182 rows and 48 columns.")
    add_body(doc,
        "02_feature_engineering converted the cleaned CSV into the "
        "model-ready feature set. Ordered categorical bands (age, "
        "education, income, job duration, working hours, working days, "
        "deliveries per day, rest break) were encoded to consecutive "
        "integers preserving order. Vehicle type and carrying mode were "
        "converted to ranks so motorbike ranks higher than scooter for "
        "vibration exposure and handheld ranks higher than backpack and "
        "bike storage for contact stress. Composite scores were computed: "
        "workload_score as the mean of the six NASA-TLX items (with the "
        "dissatisfied item reversed), fatigue_score as the mean of the "
        "six Borg items, force_exertion as the Borg lifting item, and "
        "vibration_index as vehicle_rank times work_hours_num. Five "
        "product-form interactions were added (workload times fatigue, "
        "workload times age, force times age, fatigue times job duration, "
        "deliveries times days). The eighteen NMQ and 7-day items and the "
        "five outcome follow-ups were converted to short-named binary "
        "aliases (nmq_neck, d7_lower_back, out_reduced_perf, and so on).")
    add_body(doc,
        "The posture xlsx did not carry rider identifiers, so it had to "
        "be paired with the survey without a natural join key. A "
        "severity-rank merge was used: riders were ranked by an aggregate "
        "discomfort score (NMQ pain count plus fatigue score plus daily "
        "hours), observations were ranked by RULA Table C, and the two "
        "ranked lists were matched one-to-one by position. The merged "
        "model_ready.csv contains 182 rows and 121 columns.")

    add_section_heading(doc, "3.5",
                        "Statistical analysis and model training runs")
    add_body(doc,
        "Two independent statistical analyses were run in 05_stats on the "
        "model-ready CSV. First, six chi-square tests of independence "
        "checked whether each Stage-1 risk-factor band was associated "
        "with a binary discomfort outcome; the tests reported chi2, "
        "degrees of freedom, and a Bonferroni-adjusted p-value. Second, "
        "a multivariable logistic regression fitted with statsmodels "
        "estimated odds ratios for eleven predictors (age, income, "
        "education, job duration, workload score, fatigue score, force "
        "exertion, vibration index, contact rank, hours, days) with "
        "self-reported discomfort as the outcome. Results were reported "
        "as OR with 95 percent confidence intervals.")
    add_body(doc,
        "Model training ran in 06_modeling. Every one of the seven "
        "candidate algorithms was wrapped in an imblearn.Pipeline of "
        "SMOTE followed by the classifier and hyperparameter-tuned via "
        "GridSearchCV. For each of the six risk-factor targets, this "
        "expanded to a full five-fold stratified cross-validation over "
        "the algorithm's hyperparameter grid. Across the six targets and "
        "seven algorithms, forty-two target-algorithm training runs were "
        "carried out. Best-by-macro-F1 was selected per target and the "
        "winning pipeline was refit on the full sample and persisted with "
        "joblib.dump to outputs/models/best_<factor>.pkl. The per-target "
        "feature exclusions listed in Section 2.6.9 were applied before "
        "each training run so no model could memorise its Stage-1 rule.")
    add_body(doc,
        "07_evaluation produced the reporting artefacts. cross_val_predict "
        "was used to generate out-of-fold predictions on the training "
        "set, and the following outputs were written: per-class precision, "
        "recall, F1, and support numbers via classification_report; "
        "confusion matrices per target; one-vs-rest ROC curves with "
        "per-class AUC (macro AUC is the unweighted mean of per-class "
        "AUCs); and feature importance plots (feature_importances_ for "
        "the tree-based winners, absolute value of coef_ for the Logistic "
        "Regression floor).")

    add_section_heading(doc, "3.6",
                        "Web application development and deployment")
    add_body(doc,
        "The Streamlit web application (app/streamlit_app.py) exposes a "
        "six-section form that mirrors the entire 36-item questionnaire "
        "plus the 11 RULA components and 8 QEC scores. Selectboxes cover "
        "demographics and work pattern; radios cover NMQ, 7-day, and "
        "outcome items; sliders cover NASA-TLX (0 to 100) and Borg CR10 "
        "(0 to 10); selectboxes cover RULA components; and number inputs "
        "cover the QEC scores with min / max ranges matched to the "
        "observed training data. Three sample-profile buttons (Low, "
        "Average, and High risk) fill the entire form on click so the "
        "reviewer can trigger a prediction in one step.")
    add_body(doc,
        "On form submission, encode_rider assembles the 63-feature input "
        "vector, and predict_risks runs each of the six saved models. "
        "The Results page renders a coloured metric strip (green for "
        "Low, amber for Medium, red for High), a horizontal Altair bar "
        "chart of the same six factors sorted to match the metric strip, "
        "a summary banner that flags high-burden profiles, per-factor "
        "recommendation cards, and an expandable JSON view of the exact "
        "63 feature values that were fed to the models. The app also has "
        "a Methodology page that summarises the pipeline and an About "
        "page carrying the institutional context.")
    add_body(doc,
        "The app was deployed to Streamlit Community Cloud from the "
        "public GitHub repository. Deployment required pinning every "
        "runtime dependency in requirements.txt to the exact versions "
        "the models were trained with (streamlit 1.58, pandas 2.3, "
        "scikit-learn 1.7, xgboost 3.2, imbalanced-learn 0.14, plotly "
        "6.8), because sklearn and xgboost pickles are not portable "
        "across minor versions. runtime.txt pins Python 3.13 to match "
        "the training environment. A .streamlit/config.toml carries the "
        "dark theme and a client.toolbarMode = viewer flag; CSS in the "
        "app suppresses the Fork button and 'View source on GitHub' "
        "icon that Streamlit Community Cloud injects on public-repo "
        "deploys. Every git push to the main branch auto-rebuilds and "
        "redeploys the app in about 30 seconds.")

    add_section_heading(doc, "3.7", "Iterations and fixes")
    add_subsection_heading(doc, "3.7.1", "Repetition binning correction")
    add_body(doc,
        "The initial Phase 3 binning used pandas.qcut(q=3) on "
        "deliveries_per_hour. The 66.7th percentile fell exactly on 3.889 dph "
        "(equal to 35 deliveries over 9 hours, the worst real combination in "
        "the sample). Fifty-five of the 182 riders tied at that value, and "
        "qcut is right-inclusive at the upper edge, so all 55 fell into "
        "Medium. This left the Stage-2 model unable to predict High for the "
        "worst real schedule. The fix replaces qcut with fixed cuts at 2.5 "
        "and 3.75. The Stage-1 High count grew from 19 to 74. Cross-validation "
        "accuracy dropped from 74 percent to 62 percent because the model now "
        "solves a real 3-class problem instead of memorising a 19-row minority.")

    add_subsection_heading(doc, "3.7.2", "Duration leakage correction")
    add_body(doc,
        "An earlier Phase 6 run reported 100 percent cross-validation accuracy "
        "on Duration. Investigation showed the trees were reconstructing "
        "work_hours_num from vibration_index, which equals vehicle_rank "
        "multiplied by work_hours_num. Adding vibration_index to the Duration "
        "exclusion list and capping max_depth returned Duration to a "
        "realistic 61 percent accuracy with 76 percent macro AUC.")

    add_subsection_heading(doc, "3.7.3", "Posture model with RULA and QEC inputs")
    add_body(doc,
        "The Posture target was initially trained only on the survey feature "
        "pool, which meant it could only infer posture indirectly from pain, "
        "fatigue, and hours. Adding the 11 RULA components and 8 QEC scores "
        "from the observation xlsx as explicit features moved Posture accuracy "
        "from 89 to 97 percent and the overfit gap from 0.088 to 0.028. Per "
        "the label-leakage rule, the three derived RULA Table scores were "
        "excluded because posture_risk is rula_table_c binned.")


def chapter_4(doc):
    add_chapter_heading(doc, 4, "Results and Discussions")

    add_section_heading(doc, "4.1", "Sample profile")
    add_body(doc,
        "The sample of 182 riders is 152 male and 30 female. Age skews young: "
        "78 riders are under 25, 66 in the 25-35 band, 30 in 36-45, and 8 in "
        "the 46-and-above band. Roughly half the sample (49 percent) works "
        "more than 8 hours per day. Ninety-four ride scooters and 88 ride "
        "motorbikes. The platform split is 97 Blinkit, 80 Zepto, and 5 who "
        "work both platforms.")
    add_figure(doc, ROOT / "outputs" / "figures" / "demographics.png",
               "Figure 4.1: Sample profile: age, platform, vehicle, "
               "and carrying mode.")

    add_section_heading(doc, "4.2", "NMQ pain prevalence and statistical predictors")
    add_body(doc,
        "84.6 percent of the riders reported pain in at least one body area "
        "in the last 12 months. Table 4.1 lists the prevalence per body area, "
        "sorted from highest to lowest.")
    add_table(doc,
              header=["Body area", "Prevalence"],
              rows=[
                  ["Lower back",      "61.5%"],
                  ["Upper back",      "49.5%"],
                  ["Shoulders",       "46.7%"],
                  ["Wrists / Hands",  "45.1%"],
                  ["Hips / Thighs",   "41.2%"],
                  ["Ankles / Feet",   "40.7%"],
                  ["Knees",           "39.6%"],
                  ["Neck",            "37.9%"],
                  ["Elbows",          "33.5%"],
              ],
              widths_cm=[7.0, 4.0])
    add_line(doc, "Table 4.1: NMQ 12-month pain prevalence per body area.",
             size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=12)

    add_figure(doc, ROOT / "outputs" / "figures" / "nordic_prevalence.png",
               "Figure 4.2: NMQ 12-month pain prevalence per body area.")
    add_figure(doc, ROOT / "outputs" / "figures" / "discomfort_by_demographic.png",
               "Figure 4.3: Discomfort prevalence broken down by demographic "
               "group.")

    add_body(doc,
        "The chi-square test of independence identifies Posture and Force as "
        "the two factors with a significant association with self-reported "
        "discomfort at p less than 0.05:")
    add_table(doc,
              header=["Factor", "chi-square", "p", "Significant"],
              rows=[
                  ["Posture",         "45.665", "<0.001", "Yes"],
                  ["Repetition",      "8.617",  "0.014",  "Yes"],
                  ["Force",           "6.718",  "0.035",  "Yes"],
                  ["Duration",        "0.623",  "0.733",  "No"],
                  ["Vibration",       "0.600",  "0.741",  "No"],
                  ["Contact Stress",  "0.544",  "0.762",  "No"],
              ],
              widths_cm=[4.0, 3.5, 3.0, 3.5])
    add_line(doc, "Table 4.2: Chi-square test: risk factor vs "
                  "self-reported discomfort.",
             size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=12)

    add_body(doc,
        "The multivariable logistic regression identifies seven significant "
        "individual predictors of discomfort while controlling for the others. "
        "Age is the strongest: each step up in the age band multiplies the "
        "odds of discomfort by 3.58 (95 percent CI 1.70 to 7.56, p = 0.0008).")
    add_table(doc,
              header=["Predictor", "OR", "95% CI", "p"],
              rows=[
                  ["Workload score (per unit)",     "1.06", "1.02 to 1.09", "0.0005"],
                  ["Age (per band)",                "3.58", "1.70 to 7.56", "0.0008"],
                  ["Job duration (per band)",       "2.89", "1.54 to 5.41", "0.001"],
                  ["Fatigue score (per unit)",      "1.43", "1.13 to 1.80", "0.003"],
                  ["Income (per band)",             "2.00", "1.25 to 3.20", "0.004"],
                  ["Education (per band)",          "0.33", "0.12 to 0.94", "0.04"],
                  ["Deliveries per day",            "1.05", "1.00 to 1.10", "0.045"],
              ],
              widths_cm=[6.0, 2.0, 3.5, 2.5])
    add_line(doc, "Table 4.3: Significant predictors of discomfort from "
                  "multivariable logistic regression.",
             size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=12)

    add_section_heading(doc, "4.3", "Stage-1 risk distribution")
    add_table(doc,
              header=["Factor", "Low", "Medium", "High"],
              rows=[
                  ["Force",          "90", "57",  "35"],
                  ["Repetition",     "26", "82",  "74"],
                  ["Duration",       "37", "56",  "89"],
                  ["Vibration",      "67", "68",  "47"],
                  ["Contact Stress", "68", "58",  "56"],
                  ["Posture",        "0",  "29",  "153"],
              ],
              widths_cm=[5.0, 2.5, 2.5, 2.5])
    add_line(doc, "Table 4.4: Stage-1 risk band counts per factor (n = 182).",
             size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=12)
    add_body(doc,
        "Posture, Duration, and Repetition are the three factors where the "
        "High band dominates. Posture is at 84 percent High (153 of 182 "
        "observations); Duration is at 49 percent High; Repetition is at 41 "
        "percent High after the binning fix.")
    add_figure(doc, ROOT / "outputs" / "figures" / "risk_factor_distribution.png",
               "Figure 4.4: Stage-1 Low / Medium / High counts per risk factor.")
    add_figure(doc, ROOT / "outputs" / "figures" / "risk_vs_discomfort.png",
               "Figure 4.5: Discomfort prevalence within each Low / Medium / "
               "High band, per risk factor.")
    add_figure(doc, ROOT / "outputs" / "figures" / "correlation_heatmap.png",
               "Figure 4.6: Pearson correlation matrix across the numeric "
               "feature pool, upper triangle masked.")

    add_section_heading(doc, "4.4", "Stage-2 model performance")
    add_table(doc,
              header=["Factor", "Best model", "Acc.", "F1 macro", "Macro AUC", "Features"],
              rows=[
                  ["Force",          "HistGradientBoosting", "62%", "57%", "71%", "42"],
                  ["Repetition",     "RandomForest",         "62%", "57%", "73%", "41"],
                  ["Duration",       "RandomForest",         "61%", "58%", "76%", "42"],
                  ["Vibration",      "ExtraTrees",           "58%", "57%", "72%", "41"],
                  ["Contact Stress", "RandomForest",         "60%", "59%", "74%", "42"],
                  ["Posture",        "HistGradientBoosting", "97%", "95%", "98%", "63"],
              ],
              widths_cm=[3.2, 3.8, 1.8, 2.0, 2.2, 2.0])
    add_line(doc, "Table 4.5: Best Stage-2 model per risk factor after 5-fold "
                  "stratified CV.",
             size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=12)
    add_body(doc,
        "The five survey-derived factors land between 58 and 62 percent "
        "cross-validation accuracy with macro AUC in the 71 to 76 percent "
        "range. Posture reaches 97 percent accuracy and 98 percent macro AUC "
        "because it is the only model that receives real observation inputs "
        "(the 11 RULA components and 8 QEC scores) on top of the survey "
        "features.")

    add_section_heading(doc, "4.5", "Per-class metrics and feature importance")
    add_table(doc,
              header=["Factor", "Low", "Medium", "High"],
              rows=[
                  ["Force",          "0.798", "0.580", "0.760"],
                  ["Repetition",     "0.745", "0.669", "0.782"],
                  ["Duration",       "0.822", "0.649", "0.797"],
                  ["Vibration",      "0.775", "0.615", "0.779"],
                  ["Contact Stress", "0.763", "0.651", "0.791"],
                  ["Posture",        "-",     "-",     "0.984"],
              ],
              widths_cm=[3.5, 2.5, 2.5, 2.5])
    add_line(doc, "Table 4.6: Per-class ROC AUC (one-vs-rest) for the best "
                  "model per factor.",
             size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=12)

    add_figure(doc, ROOT / "outputs" / "figures" / "confusion_matrices.png",
               "Figure 4.7: Confusion matrices for the best model per factor "
               "(rows = true class, columns = predicted class).")
    add_figure(doc, ROOT / "outputs" / "figures" / "roc_curves.png",
               "Figure 4.8: ROC curves (one-vs-rest) for the best model per "
               "factor.")

    add_body(doc,
        "Table 4.7 gives per-class precision, recall, F1, and support for the "
        "best model per factor. The Repetition High class has 65 percent "
        "precision and 70 percent recall; the Duration High class has 69 "
        "percent precision and 73 percent recall.")
    add_table(doc,
              header=["Factor", "Class", "Precision", "Recall", "F1", "Support"],
              rows=[
                  ["Force", "Low",    "0.734", "0.767", "0.750", "90"],
                  ["Force", "Medium", "0.490", "0.421", "0.453", "57"],
                  ["Force", "High",   "0.487", "0.543", "0.514", "35"],
                  ["Repetition", "Low",    "0.562", "0.346", "0.429", "26"],
                  ["Repetition", "Medium", "0.593", "0.622", "0.607", "82"],
                  ["Repetition", "High",   "0.650", "0.703", "0.675", "74"],
                  ["Duration", "Low",    "0.583", "0.757", "0.659", "37"],
                  ["Duration", "Medium", "0.450", "0.321", "0.375", "56"],
                  ["Duration", "High",   "0.691", "0.730", "0.710", "89"],
                  ["Vibration", "Low",    "0.583", "0.627", "0.604", "67"],
                  ["Vibration", "Medium", "0.551", "0.559", "0.555", "68"],
                  ["Vibration", "High",   "0.610", "0.532", "0.568", "47"],
                  ["Contact Stress", "Low",    "0.658", "0.588", "0.621", "68"],
                  ["Contact Stress", "Medium", "0.526", "0.552", "0.539", "58"],
                  ["Contact Stress", "High",   "0.617", "0.661", "0.638", "56"],
                  ["Posture", "Medium", "0.724", "0.724", "0.724", "29"],
                  ["Posture", "High",   "0.948", "0.948", "0.948", "153"],
              ],
              widths_cm=[3.5, 2.0, 2.5, 2.2, 2.0, 2.0])
    add_line(doc, "Table 4.7: Per-class precision, recall, F1, and support.",
             size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=12)

    add_figure(doc, ROOT / "outputs" / "figures" / "feature_importance.png",
               "Figure 4.9: Top 10 features by importance for the best model "
               "per factor.")
    add_table(doc,
              header=["Factor", "Rank 1", "Rank 2", "Rank 3", "Rank 4", "Rank 5"],
              rows=[
                  ["Repetition",     "income_ord",     "vibration_index",    "fatigue_score",  "workload_x_fatigue", "out_riding_worsens"],
                  ["Duration",       "deliveries_num", "deliv_x_days",       "income_ord",     "fatigue_x_jobdur",   "rest_break_num"],
                  ["Vibration",      "deliveries_num", "income_ord",         "deliv_x_days",   "rest_break_num",     "workload_score"],
                  ["Contact Stress", "vibration_index","workload_x_fatigue", "fatigue_score",  "workload_x_age",     "deliv_x_days"],
              ],
              widths_cm=[3.0, 2.8, 3.0, 2.8, 3.0, 3.0])
    add_line(doc, "Table 4.8: Top 5 most important features per factor. Force "
                  "and Posture are omitted because HistGradientBoosting does "
                  "not expose split-based importances directly.",
             size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=12)

    add_table(doc,
              header=["Factor", "Best model", "Hyperparameters", "SMOTE k"],
              rows=[
                  ["Force",          "HistGBM",      "max_depth=None, learning_rate=0.05",             "5"],
                  ["Repetition",     "RandomForest", "n_estimators=300, max_depth=None",              "5"],
                  ["Duration",       "RandomForest", "n_estimators=500, max_depth=5",                 "5"],
                  ["Vibration",      "ExtraTrees",   "n_estimators=300, max_depth=None",              "5"],
                  ["Contact Stress", "RandomForest", "n_estimators=500, max_depth=None",              "5"],
                  ["Posture",        "HistGBM",      "max_depth=5, learning_rate=0.05",               "5"],
              ],
              widths_cm=[3.2, 2.8, 6.0, 2.0])
    add_line(doc, "Table 4.9: Winning hyperparameters per target after "
                  "GridSearchCV.",
             size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=12)

    add_section_heading(doc, "4.6", "Web application demonstration")
    add_body(doc,
        "The Streamlit web application is the mentor-facing artefact of the "
        "project. Figures 4.10 through 4.14 show one full pass through the "
        "form and the resulting prediction output. The header contains three "
        "sample-profile shortcuts (Low, Average, High risk) that pre-fill the "
        "entire form so the mentor can trigger a prediction in one click.")
    add_figure(doc,
        ROOT / "outputs" / "app_screenshots" / "web_02_assessment_top.png",
        "Figure 4.10: Web app: sample-profile shortcuts and demographic "
        "section (Q1 to Q17).", width_cm=15)
    add_figure(doc,
        ROOT / "outputs" / "app_screenshots" / "web_03_assessment_nmq.png",
        "Figure 4.11: Web app: Nordic Musculoskeletal Questionnaire section "
        "(Q18 to Q24).", width_cm=15)
    add_figure(doc,
        ROOT / "outputs" / "app_screenshots" / "web_04_assessment_nasa_borg.png",
        "Figure 4.12: Web app: NASA-TLX and Borg CR10 sliders (Q25 to Q36).",
        width_cm=15)
    add_figure(doc,
        ROOT / "outputs" / "app_screenshots" / "web_05_assessment_rula_qec.png",
        "Figure 4.13: Web app: RULA (11 items) and QEC (8 scores) observation "
        "sections above the Predict button.", width_cm=15)
    add_figure(doc,
        ROOT / "outputs" / "app_screenshots" / "web_06_results.png",
        "Figure 4.14: Web app: predicted risk profile output with colour-coded "
        "level cards, radar chart, tabular result, and per-factor "
        "recommendations.", width_cm=15)


def chapter_5(doc):
    add_chapter_heading(doc, 5, "Conclusions and Extensions")

    add_body(doc,
        "This project designed, implemented, and deployed a per-rider "
        "ergonomic risk screening pipeline for last-mile quick-commerce delivery workers. The "
        "pipeline turns a 36-item self-report questionnaire and a 26-column "
        "RULA-plus-QEC observation record into six standardised risk levels "
        "spanning the classical ergonomic dimensions of Force, Repetition, "
        "Posture, Duration, Contact Stress, and Vibration. The five "
        "survey-derived factors land at 58 to 62 percent accuracy with macro "
        "AUC between 71 and 76 percent; the Posture model, which uses the "
        "observation inputs directly, reaches 97 percent accuracy and 98 "
        "percent macro AUC. The final artefact is a Streamlit web application "
        "that produces the six risk levels along with per-factor "
        "recommendations from a single form submission.")

    add_body(doc,
        "Statistically the two factors most strongly associated with "
        "self-reported discomfort in this 182-rider sample are Posture and "
        "Force. Individually, age (OR 3.58 per band), job duration (OR 2.89 "
        "per band), workload score, and fatigue score are the strongest "
        "predictors of discomfort under multivariable logistic regression. "
        "These findings match the intuition that riders in higher age bands "
        "and longer careers accumulate MSD risk that is amplified by workload "
        "and fatigue.")

    add_body(doc,
        "The main limitations of this work are the self-report nature of most "
        "input variables, the approximate per-rider linkage in the posture "
        "observations (a severity-rank merge because the observation records "
        "do not carry rider identifiers), and the small sample size (n = 182) "
        "that keeps multivariable ML variance visible. Two methodological "
        "corrections were made and openly documented during the project: the "
        "Repetition binning was changed from qcut to fixed cuts to remove a "
        "boundary tie that hid the worst real schedule, and the Duration "
        "model's feature list was adjusted to remove indirect leakage through "
        "vibration_index. Both changes cost some raw accuracy but produced "
        "more meaningful models.")

    add_body(doc,
        "Five interventions for a delivery-platform operator follow directly "
        "from the results, ordered by how many riders they would reach. First, "
        "cap daily hours (roughly half the sample exceeds 8 hours per day). "
        "Second, invest in posture-focused equipment and training (84 percent "
        "of observed postures are at RULA action level 5 or above). Third, "
        "encourage bike-storage-box carrying over handheld bags. Fourth, "
        "introduce age-targeted MSD screening for riders in the 36-and-above "
        "cohort. Fifth, tune workload management at the platform level to "
        "reduce the NASA-TLX component of the workload score, which is a "
        "significant discomfort predictor at p = 0.0005.")

    add_body(doc,
        "Future extensions of the work include a longitudinal follow-up study "
        "that would let a causal claim replace the current cross-sectional "
        "association; per-rider RULA and QEC observations to eliminate the "
        "severity-rank merge; the addition of wearable accelerometer data as "
        "a Vibration proxy replacement; and a periodic re-training schedule "
        "that lets the platform update its screening tool as the workforce "
        "changes. The two-stage design leaves room for each of these "
        "extensions without disturbing the audited Stage-1 labels.")


def add_bibliography(doc):
    # Bibliography ALSO carries a Heading 2 marker at the top of the
    # first paragraph so the STYLEREF running header switches from the
    # last chapter's name to 'BIBLIOGRAPHY'.  Uses page_break_before on
    # the heading paragraph itself instead of a separate page-break
    # paragraph -- avoids stray blank pages when Chapter 5 ends near
    # a page boundary.
    p = doc.add_paragraph(style=doc.styles["Heading 2"])
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Bibliography")
    r.font.size = Pt(28)
    r.bold = True
    r.font.color.rgb = BLACK
    set_font(r, BODY_FONT)
    refs = [
        "McAtamney L, Corlett EN (1993). RULA: a survey method for the "
        "investigation of work-related upper limb disorders. Applied "
        "Ergonomics, 24(2), 91-99.",

        "David G, Woods V, Li G, Buckle P (2008). The development of the "
        "Quick Exposure Check (QEC) for assessing exposure to risk factors "
        "for work-related musculoskeletal disorders. Applied Ergonomics, "
        "39(1), 57-69.",

        "Kuorinka I, Jonsson B, Kilbom A, Vinterberg H, Biering-Sorensen F, "
        "Andersson G, Jorgensen K (1987). Standardised Nordic questionnaires "
        "for the analysis of musculoskeletal symptoms. Applied Ergonomics, "
        "18(3), 233-237.",

        "Hart SG, Staveland LE (1988). Development of NASA-TLX (Task Load "
        "Index): Results of empirical and theoretical research. Advances in "
        "Psychology, 52, 139-183.",

        "Borg G (1998). Borg's Perceived Exertion and Pain Scales. Human "
        "Kinetics Publishers, Champaign, Illinois.",

        "Chawla NV, Bowyer KW, Hall LO, Kegelmeyer WP (2002). SMOTE: "
        "Synthetic Minority Over-sampling Technique. Journal of Artificial "
        "Intelligence Research, 16, 321-357.",

        "Breiman L (2001). Random Forests. Machine Learning, 45(1), 5-32.",

        "Geurts P, Ernst D, Wehenkel L (2006). Extremely randomized trees. "
        "Machine Learning, 63(1), 3-42.",

        "Chen T, Guestrin C (2016). XGBoost: A Scalable Tree Boosting System. "
        "Proceedings of the 22nd ACM SIGKDD International Conference on "
        "Knowledge Discovery and Data Mining, 785-794.",

        "Friedman JH (2001). Greedy function approximation: A gradient "
        "boosting machine. Annals of Statistics, 29(5), 1189-1232.",

        "Wolpert DH (1992). Stacked generalization. Neural Networks, 5(2), "
        "241-259.",

        "Pedregosa F, Varoquaux G, Gramfort A, et al. (2011). Scikit-learn: "
        "Machine Learning in Python. Journal of Machine Learning Research, "
        "12, 2825-2830.",

        "Lemaitre G, Nogueira F, Aridas CK (2017). Imbalanced-learn: A Python "
        "toolbox to tackle the curse of imbalanced datasets in machine "
        "learning. Journal of Machine Learning Research, 18(17), 1-5.",

        "Kohavi R (1995). A study of cross-validation and bootstrap for "
        "accuracy estimation and model selection. Proceedings of the 14th "
        "International Joint Conference on Artificial Intelligence, 1137-1145.",

        "Cortes C, Vapnik V (1995). Support-vector networks. Machine "
        "Learning, 20(3), 273-297.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.8)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_num = p.add_run(f"[{i}] ")
        r_num.font.size = Pt(11)
        set_font(r_num, BODY_FONT)
        r_ref = p.add_run(ref)
        r_ref.font.size = Pt(11)
        set_font(r_ref, BODY_FONT)


# ============================================================================
# Build
# ============================================================================

def build():
    doc = Document()

    # Document defaults
    s = doc.styles["Normal"]
    s.font.name = BODY_FONT
    s.font.size = Pt(11)

    # Set page margins
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.5)
        section.left_margin = section.right_margin = Cm(2.5)

    # Metadata (author + title, clear generator string)
    cp = doc.core_properties
    cp.author = STUDENT_NAME
    cp.last_modified_by = STUDENT_NAME
    cp.title = PROJECT_TITLE
    cp.comments = ""

    add_title_page(doc)

    # Section break after title page: front matter restarts at Roman 'i'
    # here so Certificate is 'i' (not 'ii').  Title page itself carries
    # no page number because section 0's footer is left empty.
    start_front_matter_section(doc)

    add_certificate(doc)
    add_abstract(doc)
    add_acknowledgements(doc)
    add_toc_placeholder(doc)
    add_list_of_figures(doc)
    add_list_of_tables(doc)
    add_abbreviations(doc)
    add_symbols(doc)

    # ---- Front matter ends here.  Section break: main body starts with
    #      its own running header (chapter STYLEREF only) and footer
    #      (PAGE only).  Page numbers restart at 1.
    start_main_body_section(doc)

    # Chapter 1 skips its internal page break -- the section break above
    # already puts us on a fresh page.
    chapter_1(doc, skip_page_break=True)
    chapter_2(doc)
    chapter_3(doc)
    chapter_4(doc)
    chapter_5(doc)
    add_bibliography(doc)

    # Make Word refresh PAGE / STYLEREF / TOC fields on open so the
    # running header shows the actual chapter name without the reader
    # having to right-click 'Update Field'.
    _enable_field_auto_update(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    p = build()
    print(f"saved {p.relative_to(ROOT)}")
