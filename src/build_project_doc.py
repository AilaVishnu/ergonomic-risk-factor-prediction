"""
Build docs/PROJECT_REPORT.docx from docs/PROJECT_REPORT.md.

Single comprehensive Word write-up of the whole project. Handles
ATX headings, bullets, fenced code, pipe tables, **bold**, *italic*,
and `inline code`.

Run:
    python src/build_project_doc.py
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SRC  = ROOT / "docs" / "PROJECT_REPORT.md"
OUT  = ROOT / "docs" / "PROJECT_REPORT.docx"

NAVY   = RGBColor(0x1F, 0x3B, 0x73)
ACCENT = RGBColor(0x2E, 0x86, 0xAB)
BLACK  = RGBColor(0x22, 0x22, 0x22)
MUTED  = RGBColor(0x55, 0x55, 0x55)


def set_font(run, font):
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), font)


def add_runs(p, text):
    """Inline parse for **bold**, *italic*, `code`."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            r = p.add_run(text[last:m.start()])
            set_font(r, "Calibri")
            r.font.size = Pt(11)
        tok = m.group(0)
        if tok.startswith("**"):
            r = p.add_run(tok[2:-2])
            set_font(r, "Calibri")
            r.font.size = Pt(11)
            r.bold = True
        elif tok.startswith("`"):
            r = p.add_run(tok[1:-1])
            set_font(r, "Consolas")
            r.font.size = Pt(10)
        else:
            r = p.add_run(tok[1:-1])
            set_font(r, "Calibri")
            r.font.size = Pt(11)
            r.italic = True
        last = m.end()
    if last < len(text):
        r = p.add_run(text[last:])
        set_font(r, "Calibri")
        r.font.size = Pt(11)


def add_heading(doc, text, level):
    sizes = {1: 18, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, "Calibri")
    r.font.size = Pt(sizes.get(level, 11))
    r.bold = True
    r.font.color.rgb = NAVY if level == 1 else ACCENT


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_runs(p, text)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    add_runs(p, text)


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.6)
        r = p.add_run(line if line else " ")
        set_font(r, "Consolas")
        r.font.size = Pt(9.5)
        r.font.color.rgb = BLACK
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table(doc, header, rows):
    n = len(header)
    t = doc.add_table(rows=1, cols=n)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        set_font(r, "Calibri")
        r.font.size = Pt(10)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], "1F3B73")
    for row in rows:
        cells = t.add_row().cells
        for i, txt in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            add_runs(p, txt)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)
    return t


def parse_table(lines, start):
    def split_row(s):
        s = s.strip()
        if s.startswith("|"): s = s[1:]
        if s.endswith("|"):   s = s[:-1]
        return [c.strip() for c in s.split("|")]

    header = split_row(lines[start])
    i = start + 2
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append(split_row(lines[i]))
        i += 1
    return header, rows, i - start


def build():
    doc = Document()

    # Document metadata - match the project owner so file properties
    # don't carry the python-docx default
    cp = doc.core_properties
    cp.author = "Vishnu vardhan"
    cp.last_modified_by = "Vishnu vardhan"
    cp.title = "Ergonomic Risk Factor Prediction - Project Report"
    cp.comments = ""

    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.0)
        section.left_margin = section.right_margin = Cm(2.0)

    # Title page header (the .md's top H1 will be skipped to avoid duplication)
    add_heading(doc, "Ergonomic Risk Factor Prediction", 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Per-rider risk profiling for food-delivery riders. "
                  "Complete project report covering background, methods, "
                  "results, web application, limitations, and reproducibility.")
    set_font(r, "Calibri")
    r.font.size = Pt(11)
    r.italic = True
    r.font.color.rgb = MUTED

    src_lines = SRC.read_text(encoding="utf-8").splitlines()
    i = 0
    n = len(src_lines)

    while i < n:
        line = src_lines[i]
        stripped = line.strip()

        # Skip the first H1 (used as title block above)
        if stripped.startswith("# ") and i < 5:
            i += 1
            continue

        if stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), 3)
            i += 1
            continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:].strip(), 2)
            i += 1
            continue
        if stripped.startswith("# "):
            add_heading(doc, stripped[2:].strip(), 1)
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            block = []
            while i < n and not src_lines[i].strip().startswith("```"):
                block.append(src_lines[i])
                i += 1
            add_code_block(doc, block)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < n and re.match(r"^\|[-:\| ]+\|?$", src_lines[i + 1].strip()):
            header, rows, consumed = parse_table(src_lines, i)
            add_table(doc, header, rows)
            i += consumed
            continue

        if stripped.startswith("- ") or re.match(r"^\d+\.\s", stripped):
            text = stripped
            if text.startswith("- "):
                text = text[2:]
            else:
                text = re.sub(r"^\d+\.\s+", "", text)
            add_bullet(doc, text)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        para_lines = [stripped]
        i += 1
        while i < n and src_lines[i].strip() and not src_lines[i].strip().startswith(("#", "-", "|", "```")):
            nxt = src_lines[i].strip()
            if re.match(r"^\d+\.\s", nxt):
                break
            para_lines.append(nxt)
            i += 1
        joined = " ".join(para_lines)
        add_paragraph(doc, joined)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"saved {path.relative_to(ROOT)}")
